import docx
from docx import enum
from docx import text
from docx.enum.text import WD_COLOR_INDEX

def combine_runs(document_name):
    document = docx.Document(document_name)
    for paragraph in document.paragraphs:
        paragraph_text = ''
        for run in paragraph.runs:
            print(paragraph.runs)
            paragraph_text += run.text
            run.text = ''
        new_run = paragraph.add_run(paragraph_text)
        # index_in_paragraph = paragraph._p.index(0)
        paragraph._p[0 + 1:0 + 1] = [new_run.element]
    document.save(document_name)
combine_runs('trbygoogle_test.docx')
doc = docx.Document('trbygoogle_test.docx')
for para in doc.paragraphs:
    if 'ВИКА' in para.text:
        for run in para.runs:
            print(run.text)
            if 'ВИКА' in run.text:
                x = run.text.split('ВИКА')
                print(x)
                print(run)
                split_index = 2
                text_before_split = run.text[0:split_index]
                text_after_split = run.text[split_index:]
                run.text = text_before_split
                newRun = para.add_run(text_after_split)
                newRun.font.highlight_color = WD_COLOR_INDEX.YELLOW
                index_in_paragraph = para._p.index(run.element)
                para._p[index_in_paragraph + 1:index_in_paragraph + 1] = [newRun.element]

doc.save('trbygoogle_test.docx')
